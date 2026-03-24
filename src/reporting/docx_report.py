"""
docx_report.py — generate report.docx from final_report.md + manual graphs.

Uses python-docx (already installed).
Parses Markdown line-by-line: headings, bullet lists, tables, paragraphs.
Inserts manual graph images after their matched section headings.
"""
import logging
import re
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .manual_graph_inserter import GraphEntry, load_manual_graphs

logger = logging.getLogger(__name__)

# Map section_key → list of markdown heading substrings that trigger graph insertion
_SECTION_TRIGGERS: Dict[str, List[str]] = {
    "role":         ["AI在制药中的角色", "（一）AI在制药中的角色"],
    "market_size":  ["市场规模与增速", "（二）市场规模与增速"],
    "sourcing_map": ["Sourcing", "公司筛选与市场地图", "2.1"],
    "competition":  ["竞争格局", "（四）竞争格局"],
    "biz_model":    ["商业模式", "（二）商业模式分类"],
}

# Reverse map: heading substring → section_key (built at module load)
_TRIGGER_TO_KEY: Dict[str, str] = {}
for _key, _triggers in _SECTION_TRIGGERS.items():
    for _t in _triggers:
        _TRIGGER_TO_KEY[_t] = _key


def _set_run_font(run, size_pt: int = 11, bold: bool = False, color: Optional[tuple] = None):
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = "微软雅黑"
    if color:
        run.font.color.rgb = RGBColor(*color)


def _add_heading(doc: Document, text: str, level: int):
    """Add a heading paragraph with correct style."""
    style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}
    style = style_map.get(level, "Heading 3")
    p = doc.add_heading(text, level=level)
    p.style = doc.styles[style]
    return p


def _add_image(doc: Document, entry: GraphEntry):
    """Insert an image with caption."""
    try:
        doc.add_picture(str(entry.path), width=Inches(5.5))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logger.info("[DocxReport] Inserted image: %s", entry.path.name)
    except Exception as exc:
        logger.warning("[DocxReport] Failed to insert image %s: %s", entry.path.name, exc)


def _parse_table(doc: Document, lines: List[str], start: int):
    """Parse a markdown table starting at lines[start], return (table_obj, end_index)."""
    rows = []
    i = start
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith("|"):
            break
        if re.match(r"^\|[-| :]+\|$", line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1

    if not rows:
        return None, start + 1

    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx < ncols:
                cell = table.cell(r_idx, c_idx)
                cell.text = cell_text
                if r_idx == 0:
                    for run in cell.paragraphs[0].runs:
                        run.font.bold = True
    return table, i


def _section_key_for_heading(heading_text: str) -> Optional[str]:
    """Return section_key if this heading matches any trigger."""
    for trigger, key in _TRIGGER_TO_KEY.items():
        if trigger in heading_text:
            return key
    return None


def generate_docx_report(
    output_dir: Path,
    md_path: Optional[Path] = None,
    graphs: Optional[Dict[str, List[GraphEntry]]] = None,
) -> Path:
    """
    Generate report.docx in output_dir.
    md_path defaults to output_dir/final_report.md.
    graphs defaults to load_manual_graphs().
    """
    md_path = md_path or (output_dir / "final_report.md")
    if not md_path.exists():
        raise FileNotFoundError(f"final_report.md not found: {md_path}")

    if graphs is None:
        graphs = load_manual_graphs()

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # Default paragraph font
    doc.styles["Normal"].font.name = "微软雅黑"
    doc.styles["Normal"].font.size = Pt(11)

    lines = md_path.read_text(encoding="utf-8").splitlines()

    # Track which section keys have been inserted already to avoid duplication
    inserted_keys: set = set()

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ── Skip horizontal rules ─────────────────────────────────────
        if stripped in ("---", "***", "___"):
            i += 1
            continue

        # ── Headings ──────────────────────────────────────────────────
        hm = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if hm:
            level = len(hm.group(1))
            text = hm.group(2).strip()
            # Strip bold markers from heading text
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            _add_heading(doc, text, level)

            # Check if we should insert graphs after this heading
            key = _section_key_for_heading(text)
            if key and key in graphs and key not in inserted_keys:
                for entry in graphs[key]:
                    _add_image(doc, entry)
                inserted_keys.add(key)

            i += 1
            continue

        # ── Tables ────────────────────────────────────────────────────
        if stripped.startswith("|"):
            _, end = _parse_table(doc, lines, i)
            i = end
            doc.add_paragraph()  # spacing after table
            continue

        # ── Bullet points ─────────────────────────────────────────────
        bm = re.match(r"^(\s*)[-*•]\s+(.*)", line)
        if bm:
            indent = len(bm.group(1))
            text = bm.group(2).strip()
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            text = re.sub(r"\[(\d+)\]", r"[\1]", text)   # keep numbered refs
            # Strip other markdown inline marks
            text = re.sub(r"`([^`]+)`", r"\1", text)
            level_style = "List Bullet 2" if indent > 0 else "List Bullet"
            p = doc.add_paragraph(style=level_style)
            p.add_run(text)
            i += 1
            continue

        # ── Blank line ────────────────────────────────────────────────
        if not stripped:
            i += 1
            continue

        # ── Regular paragraph ─────────────────────────────────────────
        text = stripped
        # Bold inline **text**
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # Split by bold markers
        parts = re.split(r"\*\*(.+?)\*\*", text)
        for idx, part in enumerate(parts):
            if not part:
                continue
            run = p.add_run(part)
            run.font.bold = (idx % 2 == 1)
            run.font.size = Pt(11)
            run.font.name = "微软雅黑"

        i += 1

    out_path = output_dir / "report.docx"
    doc.save(str(out_path))
    logger.info("[DocxReport] Saved → %s", out_path)
    return out_path
